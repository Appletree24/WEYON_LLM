import uuid
from typing import Iterable, Iterator

from docx import Document
from docx.oxml import CT_P, CT_Tbl


class Node:
    # 一个完整意义的节点，包含标签（级别），值，子节点，父节点，uuid
    def __init__(self, tag, value) -> None:
        self.tag = tag
        self.value = value
        self.children: list = []
        self.parent = None
        self.uuid: str = str(uuid.uuid4())

    def add_child(self, child):
        child.parent = self
        self.children.append(child)
        return child

    def get_value_from_tree(self):
        node = self
        parent_data = node.get_parent_values()
        text = '\n'.join(parent_data) + '\n' + node.value
        return text

    def get_parent_values(self):
        parent_data = []
        _point = self
        while _point.parent:
            if _point.parent.value:
                parent_data.append(_point.parent.value)
            _point = _point.parent
        parent_data.reverse()
        return parent_data


def convert_table_to_markdown(table):
    markdown = "|"
    for cell in table.rows[0].cells:
        markdown += f"{cell.text}|"
    markdown += "\n|"
    for _ in table.rows[0].cells:
        markdown += "---|"
    markdown += "\n"
    for row in table.rows[1:]:
        markdown += "|"
        for cell in row.cells:
            markdown += f" {cell.text} |"
        markdown += "\n"
    return markdown


class DocxLoader(Iterable[Node]):

    def __init__(self, file_path):
        self.file_path = file_path
        self.filename = file_path.split('/')[-1]
        self.document = Document(self.file_path)
        self.root = self.parse_to_tree()

    def parse_to_tree(self):
        root = Node(float('inf'), self.filename)
        doc = self.document
        point = root
        p_i = -1
        t_i = -1
        for element in doc.element.body.inner_content_elements:
            if isinstance(element, CT_P):
                p_i += 1
                para = doc.paragraphs[p_i]
                if para.runs.__len__() == 0 or para.text.strip() == '':
                    continue
                tag = para.style.font.size or 0
                while point.tag <= tag:
                    point = point.parent
                point = point.add_child(Node(tag, para.text))
            elif isinstance(element, CT_Tbl):
                t_i += 1
                table = doc.tables[t_i]
                text = convert_table_to_markdown(table)
                tag = 10.5
                _ = point.add_child(Node(tag, text))
        return root

    def __iter__(self):
        # 遍历树状结构，返回叶子节点
        def traverse(node: Node) -> Iterator[Node]:
            if node.children.__len__() == 0:
                yield node
            else:
                for child in node.children:
                    yield from traverse(child)

        yield from traverse(self.root)
