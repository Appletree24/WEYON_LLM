import os
import uuid
from typing import Iterable, Iterator

from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document


class Node:
    # 一个完整意义的节点，包含标签（级别），值，子节点，父节点，uuid
    def __init__(self, tag, value) -> None:
        self.tag = tag
        self.value = value
        self.children: list = []
        self.parent = None
        self.uuid: str = str(uuid.uuid4())

    def add_child(self, child):
        child.parent = self
        self.children.append(child)
        return child

    def get_value_from_tree(self):
        node = self
        parent_data = node.get_parent_values()
        # 从根节点开始，逐层向下遍历，遍历时在每个节点前加上n个“#”,其中n为节点级别，根节点为1
        for i, parent_datum in enumerate(parent_data):
            parent_data[i] = f"{'#' * (i + 1)} {parent_datum}"
        text = "\n".join(parent_data) + "\n" + self.value
        return text

    def get_parent_values(self):
        parent_data = []
        _point = self
        while _point.parent:
            if _point.parent.value:
                parent_data.append(_point.parent.value)
            _point = _point.parent
        parent_data.reverse()
        return parent_data


def convert_table_to_markdown(table):
    markdown = "|"
    for cell in table.rows[0].cells:
        markdown += f"{cell.text}|"
    markdown += "\n|"
    for _ in table.rows[0].cells:
        markdown += "---|"
    markdown += "\n"
    for row in table.rows[1:]:
        markdown += "|"
        for cell in row.cells:
            markdown += f" {cell.text} |"
        markdown += "\n"
    return markdown


class DocxLoader(BaseLoader, Iterable[Node]):

    def __init__(self, file_path, img_path='img', img_prefix='../img/'):
        self.img_prefix = img_prefix
        self.file_path = file_path
        self.filename = file_path.split('/')[-1]
        self.img_path = img_path
        from docx import Document
        self.document = Document(self.file_path)
        self.root = self.parse_to_tree()

    def parse_to_tree(self):
        root = Node(float('inf'), self.filename)
        doc = self.document
        paragraphs = doc.paragraphs
        point = root
        p_i = -1
        t_i = -1
        from docx.oxml import CT_P, CT_Tbl
        for element in doc.element.body.inner_content_elements:
            if isinstance(element, CT_P):
                p_i += 1
                para = paragraphs[p_i]
                if para.runs.__len__() == 0 or para.text.strip() == '':
                    images = para._element.xpath('.//pic:pic')
                    if images and paragraphs.__len__() > (p_i + 1) and paragraphs[p_i + 1].text.startswith('图'):
                        name = paragraphs[p_i + 1].text
                        tag = paragraphs[p_i + 1].style.font.size or 0
                        pre_images = self._images_handle(name, images)
                        print(name, pre_images)
                        while point.tag <= tag:
                            point = point.parent
                        point = point.add_child(Node(tag, pre_images))
                    continue
                tag = para.style.font.size or 0
                while point.tag <= tag:
                    point = point.parent
                point = point.add_child(Node(tag, para.text))
            elif isinstance(element, CT_Tbl):
                t_i += 1
                table = doc.tables[t_i]
                text = convert_table_to_markdown(table)
                point = point.parent
                point = point.add_child(Node(0, text))
        return root

    def _images_handle(self, name, images):
        res = ''
        for i, image in enumerate(images):
            image_data = image.xpath('.//a:blip/@r:embed')[0]
            related_part = self.document.part.related_parts[image_data]
            img = related_part.image
            n = f'{name}_{i}.png'
            import uuid
            filename = f'{uuid.uuid4()}.png'
            from os import path
            if not path.exists(self.img_path):
                os.makedirs(self.img_path)
            file_path = path.join(self.img_path, filename)
            fw = open(file_path, "wb")
            fw.write(img.blob)
            fw.close()
            res += f'\n![{n}]({self.img_prefix}{filename})\n'
        return res

    def __iter__(self):
        # 遍历树状结构，返回叶子节点
        def traverse(node: Node) -> Iterator[Node]:
            if node.children.__len__() == 0:
                yield node
            else:
                for child in node.children:
                    yield from traverse(child)

        yield from traverse(self.root)

    def lazy_load(self) -> Iterator[Document]:
        base_id = uuid.uuid5(uuid.uuid4(), self.file_path).int
        for node in self:
            base_id += 1
            yield Document(page_content=node.get_value_from_tree(),
                           metadata={'parent': node.parent.uuid, 'idx': base_id})
